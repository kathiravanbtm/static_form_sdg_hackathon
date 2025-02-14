import re
from flask import Flask, request, send_file, render_template
from docx import Document
import os
import io
import re 
from docx.shared import Pt 
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Function to clean the text by removing unwanted newlines and keeping paragraph separation
def replace_general_placeholders(doc, placeholders):
    """
    Replaces placeholders related to Semester, Course Name, Course Code, 
    Course Description, Prerequisites, Course Format, Assessments & Grading 
    in both paragraphs and tables while maintaining formatting.
    """
    
    # Iterate over all paragraphs
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, placeholders)

    # Iterate over all tables (tables contain rows → cells → paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, placeholders)


def replace_placeholders_in_paragraph(paragraph, placeholders):
    """
    Replaces placeholders inside a paragraph while keeping the formatting intact.
    If `<REMOVE>` is found, the paragraph is deleted.
    """
    if paragraph.runs:  # Ensure the paragraph contains text
        full_text = ''.join(run.text for run in paragraph.runs)  # Merge runs

        for placeholder, value in placeholders.items():
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, value)

        # Remove the paragraph if it contains `<REMOVE>`
        if "<REMOVE>" in full_text:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
            return

        # Apply the modified text while keeping formatting
        for i, run in enumerate(paragraph.runs):
            if i == 0:
                run.text = full_text  # Set new text in the first run
            else:
                run.text = ""  # Clear other runs




def replace_list_section(doc, placeholder, items, title=""):
    """
    Replaces a placeholder with a properly formatted numbered list while keeping the content at the correct position.
    - `placeholder`: The placeholder text to replace (e.g., `{Objectives}`)
    - `items`: The list of items to insert
    - `title`: The title of the section (optional)
    """
    for i, paragraph in enumerate(doc.paragraphs):
        if placeholder in paragraph.text:
            parent = paragraph._element.getparent()  # Get parent XML element
            paragraph.text = ""  # Clear placeholder but keep paragraph position
            if not items:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
                return 
            # Preserve the document’s original paragraph format
            paragraph_format = paragraph.paragraph_format
            
            # Insert title (if provided)
            if title:
                title_paragraph = paragraph.insert_paragraph_before()
                title_paragraph.style = paragraph.style
                title_paragraph.paragraph_format.space_before = Pt(12)  # 🔥 Space before title
                title_paragraph.paragraph_format.space_after = Pt(6)    # Keep same style
                title_run = title_paragraph.add_run(title)
                title_run.bold = True
                title_run.font.size = Pt(11)

            # Insert list items directly after the placeholder
            for index, item in enumerate(items, 1):
                item_paragraph = paragraph.insert_paragraph_before("")
                item_paragraph.style = paragraph.style  # Keep same style
                item_paragraph.paragraph_format.left_indent = paragraph_format.left_indent  # Maintain document indentation
                item_paragraph.paragraph_format.first_line_indent = paragraph_format.first_line_indent  # Keep first-line formatting
                
                # Manually add numbering (bold)
                item_run = item_paragraph.add_run(f"{index}.    ")
                item_run.bold = True  

                # Add the actual content
                content_run = item_paragraph.add_run(item.strip())  
                content_run.bold = False  
                content_run.font.size = Pt(11)

                # **🔥 Preserve Indentation & Margins using Word XML**
                pPr = item_paragraph._element.get_or_add_pPr()
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), "645")  # Use document's original left indentation
                ind.set(qn("w:hanging"), "365")  # Hanging indent for text (0.25 inch)
                pPr.append(ind)

            return

def clean_pdf_text(text):
    """Cleans extracted text while recognizing manual 'Enter' presses inside lists."""
    if not text:
        return ""

    # ✅ Trim leading/trailing spaces
    text = text.strip()

    # ✅ Fix multiple spaces and tabs
    text = re.sub(r'\s+', ' ', text)

    # ✅ Ensure correct spacing after list numbers (Fixes "1.Text" → "1. Text")
    text = re.sub(r'(\d+)\.(\S)', r'\1. \2', text)

    # ✅ Ensure proper spacing for bullet points ("-Text" → "- Text" & "•Text" → "• Text")
    text = re.sub(r'(-|•)\s*(\S)', r'\1 \2', text)

    # ✅ Preserve manual line breaks inside list items
    text = re.sub(r'(\d+\..*?)\n(\s+)(\S)', r'\1 \3', text)  # Joins lines within the same numbered item
    text = re.sub(r'(-|•)\s*(.*?)\n(\s+)(\S)', r'\1 \2 \4', text)  # Joins lines within bullet points

    return text



@app.route('/')
def index():
    return render_template('index.html')  # Load the HTML form

@app.route('/generate', methods=['POST'])
def generate_doc():
    template_path = os.path.join(BASE_DIR, "template.docx")

    if not os.path.exists(template_path):
        return "Error: template.docx not found!", 404
    doc = Document(template_path)
   # Collect form data and clean all text fields if pasted from PDF
    semester = request.form.get('Semester', '')
    course_name = request.form.get('CourseName', '')
    course_code = request.form.get('CourseCode', '')
    course_description = clean_pdf_text(request.form.get('CourseDescription', ''))
    prerequisites =clean_pdf_text(request.form.get('Prerequisites', ''))
    objectives = [clean_pdf_text(obj.strip()) for obj in request.form.getlist('objective') if obj.strip()]
    experiments = [clean_pdf_text(obj.strip()) for obj in request.form.getlist('experiments') if obj.strip()]
    course_outcomes = [clean_pdf_text(outcome.strip()) for outcome in request.form.getlist('course_outcome') if outcome.strip()]
    textbooks = [clean_pdf_text(textbook.strip()) for textbook in request.form.getlist('textbook') if textbook.strip()]
    references = [clean_pdf_text(reference.strip()) for reference in request.form.getlist('reference') if reference.strip()]
    assessments_grading = clean_pdf_text(request.form.get('AssessmentsGrading', ''))
    course_format = clean_pdf_text(request.form.get('courseformat', ''))
    assessments = clean_pdf_text(request.form.get('assessments', ''))
    grading = clean_pdf_text(request.form.get('grading', ''))
    print(course_format)
    # Format placeholders into readable lists and apply `<REMOVE>` for empty values
    placeholders = {
        "{Semester}": semester if semester else "<REMOVE>",
        "{CourseName}": course_name if course_name else "<REMOVE>",
        "{CourseCode}": course_code if course_code else "<REMOVE>",
        "{Coursedescriptionname}": "COURSE DESCRIPTION" if course_description else "<REMOVE>",
        "{CourseDescription}": course_description if course_description else "<REMOVE>",
        "{prerequisitename}": "PREREQUISITES" if prerequisites else "<REMOVE>",
        "{Prerequisites}": prerequisites if prerequisites else "<REMOVE>",
        "{AssessmentsGrading}": assessments_grading if assessments_grading else "<REMOVE>",
        "{CourseFormat}": course_format if course_format else "<REMOVE>",
        "{Assessments}": assessments if assessments else "<REMOVE>",
        "{Grading}": grading if grading else "<REMOVE>",
    }

    # Collect Practical Periods checkbox and value
    has_practical = request.form.get('hasPractical')  
    practical_periods = request.form.get('practical_periods') if has_practical else "<REMOVE>"
    # Add practical periods to placeholders if applicable
    if has_practical and practical_periods:
        placeholders["{PracticalPeriodsName}"] = "PRACTICAL PERIODS "
        placeholders["{PracticalPeriods}"] = practical_periods if practical_periods else "<REMOVE>"
    else:
        placeholders["{PracticalPeriodsName}"] = "<REMOVE>"
        placeholders["{PracticalPeriods}"] = "<REMOVE>"

    # Dynamically collect and format units including the number of periods
    units = []
    total_periods = 0
    i = 1

    while True:
        unit_title = clean_pdf_text(request.form.get(f'unit_title_{i}', ''))
        unit_content = clean_pdf_text(request.form.get(f'unit_content_{i}', ''))
        unit_periods = request.form.get(f'unit_periods_{i}')

        if not unit_title or not unit_content:
            break

        try:
            unit_periods = int(unit_periods) if unit_periods else 0
        except ValueError:
            unit_periods = 0

        total_periods += unit_periods
        units.append((unit_title, unit_content, unit_periods))
        i += 1
    # Format units into a structured text block with periods
    units_text = ""
    for i, (unit_title, unit_content, unit_periods) in enumerate(units, 1):
        units_text += f"UNIT {i}: {unit_title} (No. of Periods: {unit_periods})\n{unit_content}"
    # Add formatted units and total periods to placeholders
    placeholders["{TotalPeriods}"] ="NUMBER OF THEORY PERIODS:" + str(total_periods) if total_periods > 0 else "<REMOVE>"
    replace_list_section(doc, "{Objectives}", objectives, title="COURSE OBJECTIVES")
    replace_list_section(doc, "{Experiments}", experiments,title = "LIST OF EXPERIMENTS")
    replace_list_section(doc, "{Textbooks}", textbooks, title="TEXTBOOKS")
    replace_list_section(doc, "{References}", references, title="REFERENCES")
    format_course_outcomes(doc, "{CourseOutcomes}", course_outcomes)
    replace_units_with_formatting(doc, units)
    replace_semester(doc, semester)
    replace_course_name_in_table(doc, course_name)
    replace_course_code_in_table(doc, course_code)
    replace_course_description(doc, course_description)
    replace_prerequisites(doc, prerequisites)
    replace_course_format(doc, course_format)
    replace_assessments_grading(doc, assessments_grading)
    replace_practical_periods(doc, practical_periods)    
    total_periods = request.form.get('TotalPeriods', '')
    practical_periods = request.form.get('PracticalPeriods', '')

    # ✅ Call functions to replace placeholders
    replace_total_periods(doc, units)
    replace_practical_periods(doc, practical_periods)
    # Save and return the generated document
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(file_stream, as_attachment=True, download_name="Course_Syllabus.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def replace_semester(doc, semester):
    """Replaces the {Semester} placeholder with the actual semester value."""
    placeholder = "{Semester}"
    value = semester if semester else "<REMOVE>"

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, value)

            # Remove if marked as `<REMOVE>`
            if "<REMOVE>" in paragraph.text:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
            return


def replace_course_name_in_table(doc, course_name):
    """Finds and replaces {CourseName} inside tables while maintaining formatting."""
    placeholder = "{CourseName}"
    value = course_name if course_name else "<REMOVE>"

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        full_text = "".join(run.text for run in paragraph.runs)  # Get full text
                        new_text = full_text.replace(placeholder, value)  # Replace placeholder

                        # Clear existing runs
                        for run in paragraph.runs:
                            run.text = ""

                        # Insert new text while maintaining formatting
                        if paragraph.runs:
                            paragraph.runs[0].text = new_text
                        return  # Stop after first replacement to prevent duplicates


def replace_course_code_in_table(doc, course_code):
    """Finds and replaces {CourseCode} inside tables while maintaining formatting."""
    placeholder = "{CourseCode}"
    value = course_code if course_code else "<REMOVE>"

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        full_text = "".join(run.text for run in paragraph.runs)  # Get full text
                        new_text = full_text.replace(placeholder, value)  # Replace placeholder

                        # Clear existing runs
                        for run in paragraph.runs:
                            run.text = ""

                        # Insert new text while maintaining formatting
                        if paragraph.runs:
                            paragraph.runs[0].text = new_text
                        return  # Stop after first replacement to prevent duplicates



def replace_course_description(doc, course_description):
    """Replaces {CourseDescription} while maintaining formatting and indentation."""
    placeholder = "{CourseDescription}"
    value = course_description if course_description else "<REMOVE>"
    title = "COURSE DESCRIPTION" if course_description else "<REMOVE>"

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Preserve original paragraph formatting
            paragraph_format = paragraph.paragraph_format  # Get original indentation

            # Insert title only if a course description exists
            if course_description:
                title_paragraph = paragraph.insert_paragraph_before("")
                title_paragraph.style = paragraph.style  # Keep the same style
                title_paragraph.paragraph_format.left_indent = paragraph_format.left_indent  # Copy indentation
                title_paragraph.paragraph_format.first_line_indent = paragraph_format.first_line_indent  # Copy first-line indent

                title_run = title_paragraph.add_run(title)
                title_run.bold = True
                title_run.font.size = Pt(11)

            # Preserve formatting while replacing text
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

            # If placeholder is removed, delete paragraph
            if "<REMOVE>" in paragraph.text:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)

            return  # Stop after replacing the first occurrence



def replace_prerequisites(doc, prerequisites):
    """Adds 'PREREQUISITES' title above {Prerequisites} while maintaining formatting."""
    placeholder = "{Prerequisites}"
    title = "PREREQUISITES"
    value = prerequisites.strip() if prerequisites else "<REMOVE>"

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            p_element = paragraph._element  
            parent = p_element.getparent()

            # ✅ If prerequisites are empty, remove the placeholder paragraph
            if not prerequisites.strip():
                parent.remove(p_element)
                return  

            # ✅ Insert Title Above Without Extra Blank Paragraph
            title_paragraph = paragraph.insert_paragraph_before("")
            title_paragraph.style = paragraph.style  
            title_paragraph.paragraph_format.left_indent = paragraph.paragraph_format.left_indent  
            title_paragraph.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent  
            title_paragraph.paragraph_format.space_before = Pt(12)  # 🔥 Add space before title
            title_paragraph.paragraph_format.space_after = Pt(6)   # 🔥 Add space after title

            title_run = title_paragraph.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(11)

            # ✅ Replace placeholder with prerequisites content
            paragraph.text = value  

            # ✅ Remove the paragraph if `<REMOVE>` is present
            if "<REMOVE>" in paragraph.text:
                parent.remove(p_element)

            return  # ✅ Stop after first occurrence
  # Stop after processing the first occurrence


def replace_course_format(doc, course_format):
    """Adds 'COURSE FORMAT' title above {CourseFormat} while maintaining formatting."""
    placeholder = "{CourseFormat}"
    title = "COURSE FORMAT"
    value = course_format if course_format else "<REMOVE>"
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            if not course_format.strip():
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
                return  
            # Preserve original paragraph formatting
            paragraph_format = paragraph.paragraph_format  

            # Insert title above the placeholder
            title_paragraph = paragraph.insert_paragraph_before("")
            title_paragraph.style = paragraph.style  # Keep the same style
            title_paragraph.paragraph_format.left_indent = paragraph_format.left_indent  # Copy indentation
            title_paragraph.paragraph_format.first_line_indent = paragraph_format.first_line_indent  # Copy first-line indent

            title_run = title_paragraph.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(11)

            # Preserve the {CourseFormat} content while replacing the placeholder
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

            # Remove the paragraph if `<REMOVE>` is present
            if "<REMOVE>" in paragraph.text:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)

            return  # Stop after processing the first occurrence



def replace_assessments_grading(doc, assessments_grading):
    """Adds 'ASSESSMENTS AND GRADING' title above {AssessmentsGrading} while maintaining formatting."""
    placeholder = "{AssessmentsGrading}"
    title = "ASSESSMENTS AND GRADING"
    value = assessments_grading if assessments_grading else "<REMOVE>"
# Skip processing if there is no data

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            if not assessments_grading.strip():
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
                return  
            # Preserve original paragraph formatting
            paragraph_format = paragraph.paragraph_format  

            # Insert title above the placeholder
            title_paragraph = paragraph.insert_paragraph_before("")
            title_paragraph.style = paragraph.style  # Keep the same style
            title_paragraph.paragraph_format.left_indent = paragraph_format.left_indent  # Copy indentation
            title_paragraph.paragraph_format.first_line_indent = paragraph_format.first_line_indent  # Copy first-line indent

            title_run = title_paragraph.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(11)

            # Preserve the {AssessmentsGrading} content while replacing the placeholder
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

            # Remove the paragraph if `<REMOVE>` is present
            if "<REMOVE>" in paragraph.text:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)

            return  # Stop after processing the first occurrence



def format_objectives(doc, placeholder, objectives):
    """Replaces {Objectives} with formatted course objectives while adding a title."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            p_element = paragraph._element
            parent = p_element.getparent()
            
            if not objectives:
                parent.remove(p_element)
                return 

            paragraph.text = ""  # Clear placeholder while keeping position

            # ✅ Insert Title Above Placeholder
            title_paragraph = paragraph.insert_paragraph_before()
            title_paragraph.style = paragraph.style
            title_paragraph.paragraph_format.space_before = Pt(14)  
            title_paragraph.paragraph_format.space_after = Pt(12)
            title_paragraph.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
            title_paragraph.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent

            title_run = title_paragraph.add_run("COURSE OBJECTIVES")
            title_run.bold = True
            title_run.font.size = Pt(11)

            # ✅ Insert formatted Objectives
            for i, objective in enumerate(objectives, 1):
                obj_paragraph = paragraph.insert_paragraph_before()
                obj_paragraph.style = paragraph.style

                # **🔥 Apply Hanging Indentation using Word XML**
                pPr = obj_paragraph._element.get_or_add_pPr()
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), "950")
                ind.set(qn("w:hanging"), "740")
                pPr.append(ind)

                # First line: Numbering (bold)
                obj_run = obj_paragraph.add_run(f"{i}.   ")
                obj_run.bold = True  
                obj_run.font.size = Pt(11)

                # Content (normal font)
                content_run = obj_paragraph.add_run(objective)
                content_run.bold = False
                content_run.font.size = Pt(11)

            return


def format_textbooks(doc, placeholder, textbooks):
    """Replaces {Textbooks} with formatted textbook list while adding a title."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            p_element = paragraph._element
            parent = p_element.getparent()

            if not textbooks:
                parent.remove(p_element)
                return 

            paragraph.text = ""

            # ✅ Insert Title Above Placeholder
            title_paragraph = paragraph.insert_paragraph_before()
            title_paragraph.style = paragraph.style
            title_paragraph.paragraph_format.space_before = Pt(14)  
            title_paragraph.paragraph_format.space_after = Pt(12)

            title_run = title_paragraph.add_run("TEXTBOOKS")
            title_run.bold = True
            title_run.font.size = Pt(11)

            # ✅ Insert formatted Textbooks
            for i, textbook in enumerate(textbooks, 1):
                tb_paragraph = paragraph.insert_paragraph_before()
                tb_paragraph.style = paragraph.style

                # **🔥 Apply Hanging Indentation using Word XML**
                pPr = tb_paragraph._element.get_or_add_pPr()
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), "950")
                ind.set(qn("w:hanging"), "740")
                pPr.append(ind)

                # First line: Numbering (bold)
                tb_run = tb_paragraph.add_run(f"{i}.   ")
                tb_run.bold = True  
                tb_run.font.size = Pt(11)

                # Content (normal font)
                content_run = tb_paragraph.add_run(textbook)
                content_run.bold = False
                content_run.font.size = Pt(11)

            return


def format_references(doc, placeholder, references):
    """Replaces {References} with formatted reference list while adding a title."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            p_element = paragraph._element
            parent = p_element.getparent()

            if not references:
                parent.remove(p_element)
                return 

            paragraph.text = ""

            # ✅ Insert Title Above Placeholder
            title_paragraph = paragraph.insert_paragraph_before()
            title_paragraph.style = paragraph.style
            title_paragraph.paragraph_format.space_before = Pt(14)  
            title_paragraph.paragraph_format.space_after = Pt(12)

            title_run = title_paragraph.add_run("REFERENCES")
            title_run.bold = True
            title_run.font.size = Pt(11)

            # ✅ Insert formatted References
            for i, reference in enumerate(references, 1):
                ref_paragraph = paragraph.insert_paragraph_before()
                ref_paragraph.style = paragraph.style

                # **🔥 Apply Hanging Indentation using Word XML**
                pPr = ref_paragraph._element.get_or_add_pPr()
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), "950")
                ind.set(qn("w:hanging"), "740")
                pPr.append(ind)

                # First line: Numbering (bold)
                ref_run = ref_paragraph.add_run(f"{i}.   ")
                ref_run.bold = True  
                ref_run.font.size = Pt(11)

                # Content (normal font)
                content_run = ref_paragraph.add_run(reference)
                content_run.bold = False
                content_run.font.size = Pt(11)

            return






def format_course_outcomes(doc, placeholder, course_outcomes):
    """Replaces {CourseOutcomes} with formatted course outcomes while adding a title."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            p_element = paragraph._element
            parent = p_element.getparent()
            
            if not course_outcomes:
                parent.remove(p_element)
                return  

            paragraph.text = ""  # Clear placeholder while keeping position

            # Preserve paragraph formatting
            paragraph_format = paragraph.paragraph_format

            # ✅ Insert Title Above Placeholder
            title_paragraph = paragraph.insert_paragraph_before()
            title_paragraph.style = paragraph.style
            title_paragraph.paragraph_format.space_before = Pt(14)  
            title_paragraph.paragraph_format.space_after = Pt(12)  
            title_paragraph.paragraph_format.left_indent = paragraph_format.left_indent  
            title_paragraph.paragraph_format.first_line_indent = paragraph_format.first_line_indent  

            title_run = title_paragraph.add_run("COURSE OUTCOMES")
            title_run.bold = True
            title_run.font.size = Pt(11)

            # ✅ Insert formatted COs **without adding an empty paragraph**
            for i, outcome in enumerate(course_outcomes, 1):
                co_paragraph = paragraph.insert_paragraph_before()  # ✅ Fix: No empty string inserted
                co_paragraph.style = paragraph.style

                # **🔥 Apply Hanging Indentation using Word XML**
                pPr = co_paragraph._element.get_or_add_pPr()
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), "950")  
                ind.set(qn("w:hanging"), "740")  
                pPr.append(ind)

                # First line: CO label (bold)
                co_run = co_paragraph.add_run(f"CO{i}      ") 
                co_run.bold = True  
                co_run.font.size = Pt(11)

                # Content (normal font)
                content_run = co_paragraph.add_run(outcome)
                content_run.bold = False
                content_run.font.size = Pt(11)

            return  
        

def replace_units_with_formatting(doc, units):
    """Finds {Units} placeholder and inserts formatted units with proper indentation & normal content formatting."""
    for paragraph in doc.paragraphs:
        if "{Units}" in paragraph.text:
            p_element = paragraph._element  # Store reference to remove placeholder
            parent = p_element.getparent()  # Get parent XML element
            paragraph_style = paragraph.style  # Store the style of the original paragraph
            # Create a new paragraph at the same location before removing {Units}
            new_paragraph = paragraph.insert_paragraph_before("")
            new_paragraph.style = paragraph_style  # Apply the same style as the placeholder
            parent.remove(p_element)  # Remove {Units} placeholder

            for i, (unit_title, unit_content, unit_periods) in enumerate(units, 1):
                # Insert Unit Title (Bold) with correct style
                title_paragraph = new_paragraph.insert_paragraph_before("")
                title_paragraph.style = paragraph_style  # Apply same style
                title_run = title_paragraph.add_run(f"UNIT {i}: {unit_title} (No. of Periods: {unit_periods})")
                title_paragraph.paragraph_format.space_before = Pt(12)  # 🔥 Space before title
                title_paragraph.paragraph_format.space_after = Pt(10)  
                title_run.bold = True
                title_run.font.size = Pt(11)

                # Insert Unit Content (Normal) with correct indentation
                content_paragraph = new_paragraph.insert_paragraph_before("")
                content_paragraph.style = paragraph_style  # Apply same style  
                content_run = content_paragraph.add_run(f"{unit_content}")
                content_run.bold = False  # 🔥 Fix: Ensure normal text
                content_run.font.size = Pt(11)

            break 

def replace_practical_periods(doc, practical_periods):
    """Replaces {PracticalPeriods} with a single-line format while maintaining formatting."""
    placeholder = "{PracticalPeriods}"
    value = f"PRACTICAL PERIODS: {practical_periods}" if practical_periods else "<REMOVE>"

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # ✅ Replace the placeholder with formatted single-line text
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

            # ✅ Remove placeholder if there is no practical period
            if "<REMOVE>" in paragraph.text:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)

            return  # ✅ Stop after first occurrence


def replace_total_periods(doc, units):
    """Calculates total periods from all units and replaces {TotalPeriods} in a single line."""
    placeholder = "{TotalPeriods}"

    # ✅ Calculate total periods by summing unit periods
    total_periods = sum(unit[2] for unit in units) if units else 0
    value = f"TOTAL NUMBER OF PERIODS: {total_periods}" if total_periods > 0 else "<REMOVE>"

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # ✅ Combine all runs text (handles cases where {TotalPeriods} is split across runs)
            full_text = "".join(run.text for run in paragraph.runs)
            updated_text = full_text.replace(placeholder, value)

            # ✅ Clear existing runs before inserting updated text
            for run in paragraph.runs:
                run.text = ""

            # ✅ Set the new text in the first run
            if paragraph.runs:
                paragraph.runs[0].text = updated_text

            # ✅ Remove paragraph if `<REMOVE>` is present
            if "<REMOVE>" in paragraph.text:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)

            return  # ✅ Stop after first occurrence
def replace_list_of_experiments(doc, placeholder, experiments):
    """
    Replaces {ListOfExperiments} with a properly formatted numbered list while keeping formatting.
    - `placeholder`: The placeholder text to replace (e.g., {ListOfExperiments})
    - `experiments`: The list of experiments to insert
    """
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            parent = paragraph._element.getparent()

            # ✅ If no experiments exist, remove placeholder and exit
            if not experiments:
                parent.remove(paragraph._element)
                return  

            paragraph.text = ""  # Clear placeholder while keeping position

            # Preserve paragraph formatting
            paragraph_format = paragraph.paragraph_format
            
            # ✅ Insert Title: "PRACTICAL EXERCISES"
            title_paragraph = paragraph.insert_paragraph_before("")
            title_paragraph.style = paragraph.style  
            title_paragraph.paragraph_format.left_indent = paragraph_format.left_indent  
            title_paragraph.paragraph_format.first_line_indent = paragraph_format.first_line_indent  

            title_run = title_paragraph.add_run("PRACTICAL EXERCISES")
            title_run.bold = True
            title_run.font.size = Pt(11)

            # ✅ Insert formatted list of experiments
            for index, experiment in enumerate(experiments, 1):
                exp_paragraph = paragraph.insert_paragraph_before("")
                exp_paragraph.style = paragraph.style  
                exp_paragraph.paragraph_format.left_indent = paragraph_format.left_indent  
                exp_paragraph.paragraph_format.first_line_indent = paragraph_format.first_line_indent  

                # Add numbering (bold)
                exp_run = exp_paragraph.add_run(f"{index}. ")
                exp_run.bold = True  

                # Add the actual content
                content_run = exp_paragraph.add_run(experiment.strip())  
                content_run.bold = False  
                content_run.font.size = Pt(11)

            return  # ✅ Stop after first occurrence

if __name__ == '__main__':
    app.run(debug=True)
